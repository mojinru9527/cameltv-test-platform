"""File parser service — extract text from uploaded requirement documents."""
from __future__ import annotations

import io


def parse_markdown(file_bytes: bytes) -> str:
    """Parse .md file — read as UTF-8 text."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


def parse_docx(file_bytes: bytes) -> str:
    """Parse .docx file — extract all paragraph and table text."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


# Columns that indicate a test-case Excel sheet
_TESTCASE_COLUMNS = {"用例标题", "模块", "优先级", "重要程度", "操作步骤", "预期结果", "用例编号"}


def parse_xlsx(file_bytes: bytes) -> dict:
    """Parse .xlsx file — dual mode: requirement text or structured test cases.

    Returns {"type": "requirement" | "test_cases", "content": str, "cases": list[dict] | None}
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    ws = wb.active
    if ws is None:
        return {"type": "requirement", "content": "", "cases": None}

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return {"type": "requirement", "content": "", "cases": None}

    header = [str(c).strip() if c else "" for c in rows[0]]
    header_set = set(header)

    # Detect test-case format
    if _TESTCASE_COLUMNS & header_set:
        # Structured test case table
        cases: list[dict] = []
        for row in rows[1:]:
            values = [str(c).strip() if c is not None else "" for c in row]
            if not any(values):
                continue
            case = dict(zip(header, values))
            cases.append(case)

        content = f"Excel 测试用例表格，共 {len(cases)} 条用例。表头: {', '.join(header)}"
        return {"type": "test_cases", "content": content, "cases": cases}

    # Free-form requirement document
    lines: list[str] = []
    for row in rows:
        cells = [str(c).strip() if c is not None else "" for c in row]
        line = " | ".join(filter(None, cells))
        if line:
            lines.append(line)

    content = "\n".join(lines)
    return {"type": "requirement", "content": content, "cases": None}
