"""证据包 Word 导出 —— 人审用 .docx（截图 + 识别文本）。

Word 是人审工件，非唯一真源；每页一章，含路径、质量状态、截图与合并文本。
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Inches


@dataclass
class WordPage:
    page_name: str
    page_path: str
    merged_text: str
    quality: dict
    screenshots: list[Path] = field(default_factory=list)


def export_word(output_path: Path, title: str, source_url: str, pages: list[WordPage]) -> Path:
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"来源链接：{source_url}")
    doc.add_paragraph(f"页面数量：{len(pages)}")

    for idx, page in enumerate(pages, start=1):
        doc.add_page_break()
        doc.add_heading(f"{idx}. {page.page_name}", level=1)
        doc.add_paragraph(f"路径：{page.page_path}")
        doc.add_paragraph(f"质量状态：{page.quality.get('status', '')}")
        for shot in page.screenshots:
            shot = Path(shot)
            if shot.exists():
                doc.add_paragraph(f"截图：{shot.name}")
                try:
                    doc.add_picture(str(shot), width=Inches(6.5))
                except Exception:  # noqa: BLE001 — 单张截图损坏不应中断整份文档
                    doc.add_paragraph(f"（截图嵌入失败：{shot.name}）")
        doc.add_heading("识别文本", level=2)
        for line in (page.merged_text or "").splitlines():
            doc.add_paragraph(line)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
